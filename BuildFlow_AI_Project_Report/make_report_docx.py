from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\mini project\version 4\BuildFlow_AI_Project_Report")
INPUT_MD = ROOT / "BuildFlow_AI_Project_Report.md"
OUTPUT_DOCX = ROOT / "BuildFlow_AI_Project_Report_Final.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def style_run(run, *, bold=False, italic=False, size=12, name="Times New Roman"):
    run.bold = bold
    run.italic = italic
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(140)
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("BuildFlow AI - AI-Based Project Planning Assistant")
    style_run(r, bold=True, size=20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle, after=12, line=1.2)
    r = subtitle.add_run("Academic Project Report")
    style_run(r, bold=True, size=14)

    for text in [
        "Department of Computer Science Engineering",
        "Rajiv Gandhi College of Engineering and Technology",
        "Academic Year: 2025-2026",
        "",
        "Team Members",
        "VIDYASAAGAR M (23TD0091) - Lead Developer (Full Stack + AI Integration)",
        "AJAY M (23TD0003) - Documentation & Testing Support",
        "SHAFI AHMED F (23TD0076) - Research & Assistance",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, after=8, line=1.15)
        r = p.add_run(text)
        style_run(r, bold=(text == "Team Members"), size=12 if text else 6)

    doc.add_page_break()


def parse_blocks(text):
    lines = text.splitlines()
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            blocks.append(("code", lang, "\n".join(code_lines).rstrip()))
            i += 1
            continue

        if re.match(r"^#{1,3}\s+", stripped):
            level = len(stripped.split(" ")[0])
            content = stripped[level:].strip()
            blocks.append(("heading", level, content))
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            blocks.append(("table", table_lines))
            continue

        if re.match(r"^\d+\.\s+", stripped):
            items = []
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                items.append(re.sub(r"^\d+\.\s+", "", lines[i].strip()))
                i += 1
            blocks.append(("olist", items))
            continue

        if stripped.startswith("- "):
            items = []
            while i < len(lines) and lines[i].strip().startswith("- "):
                items.append(lines[i].strip()[2:].strip())
                i += 1
            blocks.append(("ulist", items))
            continue

        para_lines = [stripped]
        i += 1
        while i < len(lines):
            next_stripped = lines[i].strip()
            if (
                not next_stripped
                or next_stripped.startswith("```")
                or next_stripped.startswith("|")
                or next_stripped.startswith("- ")
                or re.match(r"^\d+\.\s+", next_stripped)
                or re.match(r"^#{1,3}\s+", next_stripped)
            ):
                break
            para_lines.append(next_stripped)
            i += 1
        blocks.append(("paragraph", " ".join(para_lines)))
    return blocks


def add_heading(doc, level, text):
    if level == 1 and text == "BuildFlow AI - AI-Based Project Planning Assistant":
        return
    paragraph = doc.add_paragraph(style=f"Heading {min(level, 3)}")
    set_paragraph_spacing(paragraph, before=10 if level == 1 else 6, after=6, line=1.1)
    run = paragraph.add_run(text)
    style_run(run, bold=True, size=16 if level == 1 else 14 if level == 2 else 12)


def add_paragraph_block(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, after=8, line=1.2)
    run = p.add_run(text)
    style_run(run)
    p.paragraph_format.first_line_indent = Inches(0.3)


def add_list(doc, items, ordered=False):
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph(style="Normal")
        set_paragraph_spacing(p, after=4, line=1.15)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        prefix = f"{index}. " if ordered else u"\u2022 "
        r = p.add_run(prefix + item)
        style_run(r)


def add_code_block(doc, lang, code):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=4, after=8, line=1.05)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.2)
    for line in code.splitlines() or [""]:
        run = p.add_run(line)
        style_run(run, size=10, name="Courier New")
        run.add_break()


def split_row(row):
    parts = [part.strip() for part in row.strip().strip("|").split("|")]
    return parts


def add_table(doc, table_lines):
    rows = [split_row(line) for line in table_lines if set(line.strip()) != {"|", "-"}]
    if len(rows) < 2:
        return

    header = rows[0]
    body = rows[1:]
    cols = len(header)

    table = doc.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True

    hdr_cells = table.rows[0].cells
    for idx, value in enumerate(header):
        cell = hdr_cells[idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, after=2, line=1.0)
        r = p.add_run(value)
        style_run(r, bold=True, size=11)
        set_cell_shading(cell, "FFFFFF")

    for row_values in body:
        row_cells = table.add_row().cells
        for idx in range(cols):
            cell = row_cells[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            value = row_values[idx] if idx < len(row_values) else ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(value) < 20 else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(p, after=2, line=1.05)
            r = p.add_run(value)
            style_run(r, size=10.5)

    doc.add_paragraph()


def build_docx():
    markdown = INPUT_MD.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    add_title_page(doc)

    blocks = parse_blocks(markdown)
    skip_first_heading = True
    for block in blocks:
        kind = block[0]
        if kind == "heading":
            _, level, text = block
            if skip_first_heading and level == 1 and text == "BuildFlow AI - AI-Based Project Planning Assistant":
                skip_first_heading = False
                continue
            add_heading(doc, level, text)
        elif kind == "paragraph":
            add_paragraph_block(doc, block[1])
        elif kind == "olist":
            add_list(doc, block[1], ordered=True)
        elif kind == "ulist":
            add_list(doc, block[1], ordered=False)
        elif kind == "code":
            add_code_block(doc, block[1], block[2])
        elif kind == "table":
            add_table(doc, block[1])

    doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    build_docx()
